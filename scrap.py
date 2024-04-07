from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
import time
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
from sumy.nlp.stemmers import Stemmer

import nltk
nltk.download('punkt')


# Function to fetch terms and conditions from a website
def fetch_terms_and_conditions(url):
    # Setup Selenium
    options = Options()
    options.headless = True
    service = Service(ChromeDriverManager().install())
    driver = webdriver.Chrome(service=service, options=options)
    
    # Navigate to the website
    driver.get(url)
    
    # Wait for a few seconds for the website to load completely
    time.sleep(5)  # Adjust the sleep time as needed
    
    # Find the terms and conditions element
    terms_element = driver.find_element(By.ID, 'page-terms')
    
    # Extract the text of the terms and conditions
    terms_text = terms_element.text
    
    # Close the browser
    driver.quit()
    
    return terms_text

# URL of the website containing terms and conditions
website_url = 'https://brand.netflix.com/en/terms/'

# Fetch terms and conditions from the website
terms_text = fetch_terms_and_conditions(website_url)

if terms_text:
    # Create a new Word document for summarized text
    doc_summary = Document()
    
    # Initialize the parser and tokenizer
    parser = PlaintextParser.from_string(terms_text, Tokenizer("english"))

    # Initialize the summarizers
    lsa_summarizer = LsaSummarizer(Stemmer("english"))

    # Extract summaries
    lsa_summary = lsa_summarizer(parser.document, 4)  # Number of sentences in the summary

    # Add the summarized text to the document
    doc_summary.add_heading('Summarized Terms and Conditions', level=1)
    for sentence in lsa_summary:
        doc_summary.add_paragraph(str(sentence))
    
    # Save the summarized text to a new Word document
    doc_summary_file_path = 'summarized_terms_and_conditions.docx'
    doc_summary.save(doc_summary_file_path)
    
    print(f"Summarized terms and conditions saved in the document: {doc_summary_file_path}")
else:
    print("No terms and conditions fetched from the website.")
